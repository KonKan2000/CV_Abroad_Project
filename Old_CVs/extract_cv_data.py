from docx import Document
import os

output = []
docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]

for file in docx_files:
    try:
        doc = Document(file)
        output.append(f"\n{'='*80}\nFROM FILE: {file}\n{'='*80}\n")
        for para in doc.paragraphs:
            if para.text.strip():
                output.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        output.append(cell.text.strip())
    except Exception as e:
        output.append(f"ERROR reading {file}: {e}")

with open('CV_EXTRACTED_DATA.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))

print("Extracted text saved to CV_EXTRACTED_DATA.txt")
print(f"Processed {len(docx_files)} DOCX files")
