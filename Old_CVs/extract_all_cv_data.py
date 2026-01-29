from docx import Document
import pdfplumber
import os

output = []

# Extract from DOCX files
docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
for file in docx_files:
    try:
        doc = Document(file)
        output.append(f"\n{'='*80}\nFROM FILE: {file}\n{'='*80}\n")
        for para in doc.paragraphs:
            if para.text.strip():
                output.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        output.append(cell.text.strip())
    except Exception as e:
        output.append(f"ERROR reading {file}: {e}")

# Extract from DOC files (try with docx first, older .doc format)
doc_files = [f for f in os.listdir('.') if f.endswith('.doc')]
for file in doc_files:
    try:
        doc = Document(file)
        output.append(f"\n{'='*80}\nFROM FILE: {file}\n{'='*80}\n")
        for para in doc.paragraphs:
            if para.text.strip():
                output.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        output.append(cell.text.strip())
    except Exception as e:
        output.append(f"ERROR reading {file}: {e}")

# Extract from PDF files
pdf_files = [f for f in os.listdir('.') if f.endswith('.pdf')]
for file in pdf_files:
    try:
        output.append(f"\n{'='*80}\nFROM FILE: {file}\n{'='*80}\n")
        with pdfplumber.open(file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    output.append(f"--- Page {page_num + 1} ---\n{text}")
    except Exception as e:
        output.append(f"ERROR reading {file}: {e}")

with open('CV_EXTRACTED_DATA_ALL.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))

print("Extracted text saved to CV_EXTRACTED_DATA_ALL.txt")
print(f"Processed {len(docx_files)} DOCX files, {len(doc_files)} DOC files, {len(pdf_files)} PDF files")
